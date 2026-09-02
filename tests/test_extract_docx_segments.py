import tempfile
import unittest
from pathlib import Path
import sys

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))
from extract_docx_segments import extract_docx_segments  # noqa: E402


class ExtractDocxSegmentsTests(unittest.TestCase):
    def test_extracts_sentence_units_and_structural_locations(self):
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "segments.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("Первое предложение. ")
            paragraph.add_run("Второе предложение!").bold = True
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Текст внутри таблицы"
            document.sections[0].header.paragraphs[0].text = "Текст колонтитула"
            document.sections[0].footer.paragraphs[0].text = "Номер документа"
            document.save(source)

            result = extract_docx_segments(source)

        sources = [segment["source"] for segment in result["segments"]]
        self.assertIn("Первое предложение.", sources)
        self.assertIn("Второе предложение!", sources)
        self.assertIn("Текст внутри таблицы", sources)
        self.assertIn("Текст колонтитула", sources)
        self.assertIn("Номер документа", sources)

        second = next(segment for segment in result["segments"] if segment["source"] == "Второе предложение!")
        self.assertEqual(second["kind"], "paragraph")
        self.assertTrue(second["formatting"]["allBold"])
        self.assertGreater(second["location"]["charStart"], 0)

        table_segment = next(segment for segment in result["segments"] if segment["kind"] == "table-cell")
        self.assertEqual(table_segment["location"]["tableIndex"], 0)
        self.assertEqual(table_segment["location"]["rowIndex"], 0)
        self.assertEqual(table_segment["location"]["cellIndex"], 0)

    def test_ids_are_deterministic_for_an_unchanged_document(self):
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "stable.docx"
            document = Document()
            document.add_paragraph("Один и тот же текст.")
            document.save(source)

            first = extract_docx_segments(source)
            second = extract_docx_segments(source)

            edited = Document(source)
            edited.paragraphs[0].text = "Текст был изменён."
            edited.save(source)
            after_text_edit = extract_docx_segments(source)

        self.assertEqual(
            [segment["id"] for segment in first["segments"]],
            [segment["id"] for segment in second["segments"]],
        )
        self.assertEqual(first["segments"][0]["id"], after_text_edit["segments"][0]["id"])


if __name__ == "__main__":
    unittest.main()
