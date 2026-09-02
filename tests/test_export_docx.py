import tempfile
import unittest
from pathlib import Path
import sys

from docx import Document
sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))
from export_docx import export_layout  # noqa: E402


class ExportDocxTests(unittest.TestCase):
    def test_exports_tabs_and_run_level_bold_italic_highlight(self):
        payload = {
            "title": "Rich text",
            "pages": [{"id": "page-1", "index": 0, "widthPx": 800, "heightPx": 1100}],
            "segments": [{
                "id": "rich", "pageIndex": 0, "text": "/Подпись/\t/Печать/", "x": 32, "y": 48,
                "width": 700, "height": 60, "fontFamily": "Times New Roman", "fontSizePx": 16,
                "fontWeight": 400, "fontStyle": "normal", "lineHeight": 1.2, "color": "#111827",
                "alignment": "left", "zIndex": 1, "rotation": 12,
                "runs": [
                    {"text": "/Подпись/\t", "fontWeight": 400, "fontStyle": "italic",
                     "tabStopPx": 180},
                    {"text": "/Печать/", "fontFamily": "Arial", "fontSizePx": 20,
                     "fontWeight": 700, "fontStyle": "italic", "color": "#224466",
                     "backgroundColor": "#FFFF00"},
                ],
            }],
        }

        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "rich.docx"
            export_layout(payload, output)
            document = Document(output)

        text_box = document.element.body.xpath(".//w:txbxContent")[0]
        shape = document.element.body.xpath(
            ".//*[local-name()='shape' and namespace-uri()='urn:schemas-microsoft-com:vml']"
        )[0]
        self.assertIn("rotation:12", shape.get("style"))
        self.assertEqual(len(text_box.xpath(
            ".//*[local-name()='r']/*[local-name()='tab']"
        )), 1)
        self.assertEqual(
            text_box.xpath(".//*[local-name()='pPr']/*[local-name()='tabs']/*[local-name()='tab']/@*[local-name()='pos']"),
            ["2700"],
        )
        bold_text = "".join(
            node.text or "" for node in text_box.xpath(
                ".//*[local-name()='r'][./*[local-name()='rPr']/*[local-name()='b']]/*[local-name()='t']"
            )
        )
        self.assertEqual(bold_text, "/Печать/")
        self.assertEqual(len(text_box.xpath(
            ".//*[local-name()='r'][./*[local-name()='rPr']/*[local-name()='i']]"
        )), 2)
        self.assertEqual(
            text_box.xpath(".//*[local-name()='rPr']/*[local-name()='shd']/@*[local-name()='fill']"),
            ["FFFF00"],
        )
        styled_run = text_box.xpath(
            ".//*[local-name()='r'][./*[local-name()='t' and text()='/Печать/']]"
        )[0]
        self.assertEqual(
            styled_run.xpath("./*[local-name()='rPr']/*[local-name()='rFonts']/@*[local-name()='ascii']"),
            ["Arial"],
        )
        self.assertEqual(
            styled_run.xpath("./*[local-name()='rPr']/*[local-name()='sz']/@*[local-name()='val']"),
            ["30"],
        )
        self.assertEqual(
            styled_run.xpath("./*[local-name()='rPr']/*[local-name()='color']/@*[local-name()='val']"),
            ["224466"],
        )

    def test_css_justify_is_exported_as_word_both_alignment(self):
        payload = {
            "title": "Justified text regression",
            "pages": [
                {"id": "page-1", "index": 0, "widthPx": 794, "heightPx": 1123},
                {"id": "page-2", "index": 1, "widthPx": 794, "heightPx": 1123},
            ],
            "segments": [],
        }
        for index in range(20):
            payload["segments"].append({
                "id": f"segment-{index}",
                "pageIndex": index // 10,
                "text": "ИМЕЕТ СРОК: Настоящая доверенность действительна до окончания рабочего дня.",
                "x": 48,
                "y": 32 + (index % 10) * 80,
                "width": 690,
                "height": 64,
                "fontFamily": "Times New Roman",
                "fontSizePx": 16,
                "fontWeight": 700,
                "lineHeight": 1.1,
                "color": "#000000",
                "alignment": "justify",
                "zIndex": index + 1,
            })

        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "justified.docx"
            export_layout(payload, output)
            document = Document(output)

        alignments = document.element.body.xpath(".//w:txbxContent/w:p/w:pPr/w:jc/@w:val")
        self.assertEqual(alignments, ["both"] * 20)
        self.assertNotIn("justify", alignments)

    def test_exports_positioned_editable_text_on_two_pages(self):
        payload = {
            "title": "Position test",
            "pages": [
                {"id": "page-1", "index": 0, "widthPx": 800, "heightPx": 1100},
                {"id": "page-2", "index": 1, "widthPx": 800, "heightPx": 1100},
            ],
            "segments": [
                {
                    "id": "one", "pageIndex": 0, "text": "First page", "x": 32, "y": 48,
                    "width": 300, "height": 60, "fontFamily": "Arial", "fontSizePx": 16,
                    "fontWeight": 400, "lineHeight": 1.2, "color": "#111827",
                    "alignment": "left", "zIndex": 1,
                },
                {
                    "id": "two", "pageIndex": 1, "text": "Second page", "x": 64, "y": 96,
                    "width": 320, "height": 70, "fontFamily": "Arial", "fontSizePx": 18,
                    "fontWeight": 700, "lineHeight": 1.2, "color": "#223344",
                    "alignment": "center", "zIndex": 2,
                },
            ],
        }

        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "output.docx"
            export_layout(payload, output)
            document = Document(output)

        self.assertEqual(len(document.sections), 2)
        text = "".join(node.text or "" for node in document.element.body.xpath(".//w:txbxContent//w:t"))
        self.assertIn("First page", text)
        self.assertIn("Second page", text)

        anchors = document.element.body.xpath(".//wp:anchor")
        self.assertEqual(len(anchors), 0)
        self.assertEqual(len(document.element.body.xpath(".//w:framePr")), 0)
        shapes = document.element.body.xpath(
            ".//*[local-name()='shape' and namespace-uri()='urn:schemas-microsoft-com:vml']"
        )
        self.assertEqual(len(shapes), 2)
        shape_types = document.element.body.xpath(
            ".//*[local-name()='shapetype' and namespace-uri()='urn:schemas-microsoft-com:vml']"
        )
        self.assertEqual(len(shape_types), 1)
        shape_ids = [shape.get("{urn:schemas-microsoft-com:office:office}spid") for shape in shapes]
        self.assertEqual(len(shape_ids), len(set(shape_ids)))
        shape_defaults = document.settings.element.xpath(
            ".//*[local-name()='shapedefaults' and namespace-uri()='urn:schemas-microsoft-com:office:office']"
        )
        self.assertGreater(
            max(int(node.get("spidmax", 0)) for node in shape_defaults),
            max(int(shape_id.rsplit("s", 1)[-1]) for shape_id in shape_ids),
        )
        self.assertIn("margin-left:24pt", shapes[0].get("style"))
        self.assertIn("margin-top:36pt", shapes[0].get("style"))
        self.assertIn("width:225pt", shapes[0].get("style"))
        self.assertIn("height:45pt", shapes[0].get("style"))
        self.assertIn("mso-position-horizontal-relative:page", shapes[0].get("style"))
        self.assertIn("mso-position-vertical-relative:page", shapes[0].get("style"))
        alternate_content = document.element.body.xpath(
            ".//*[local-name()='AlternateContent' and namespace-uri()='http://schemas.openxmlformats.org/markup-compatibility/2006']"
        )
        self.assertEqual(len(alternate_content), 0)

    def test_preserves_an_inserted_empty_page_between_content_pages(self):
        payload = {
            "title": "Blank middle page",
            "pages": [
                {"id": "page-1", "index": 0, "widthPx": 800, "heightPx": 1100},
                {"id": "page-2", "index": 1, "widthPx": 800, "heightPx": 1100},
                {"id": "page-3", "index": 2, "widthPx": 800, "heightPx": 1100},
            ],
            "segments": [
                {
                    "id": "first", "pageIndex": 0, "text": "Before blank page", "x": 32, "y": 48,
                    "width": 300, "height": 60, "fontFamily": "Arial", "fontSizePx": 16,
                    "fontWeight": 400, "lineHeight": 1.2, "color": "#111827",
                    "alignment": "left", "zIndex": 1,
                },
                {
                    "id": "third", "pageIndex": 2, "text": "After blank page", "x": 64, "y": 96,
                    "width": 320, "height": 70, "fontFamily": "Arial", "fontSizePx": 18,
                    "fontWeight": 700, "lineHeight": 1.2, "color": "#223344",
                    "alignment": "center", "zIndex": 2,
                },
            ],
        }

        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "blank-middle.docx"
            export_layout(payload, output)
            document = Document(output)

        self.assertEqual(len(document.sections), 3)
        shapes = document.element.body.xpath(
            ".//*[local-name()='shape' and namespace-uri()='urn:schemas-microsoft-com:vml']"
        )
        self.assertEqual(len(shapes), 2)

    def test_many_segments_use_independent_flow_anchors(self):
        pages = [
            {"id": "page-1", "index": 0, "widthPx": 800, "heightPx": 1100},
            {"id": "page-2", "index": 1, "widthPx": 800, "heightPx": 1100},
        ]
        segments = []
        for index in range(240):
            page_index = index // 120
            segments.append({
                "id": f"segment-{index}",
                "pageIndex": page_index,
                "text": f"Text {index}",
                "x": (index % 6) * 120,
                "y": ((index // 6) % 20) * 50,
                "width": 110,
                "height": 40,
                "fontFamily": "Arial",
                "fontSizePx": 12,
                "fontWeight": 400,
                "lineHeight": 1.2,
                "color": "#111827",
                "alignment": "left",
                "zIndex": index + 1,
            })

        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "many.docx"
            export_layout({"title": "Many", "pages": pages, "segments": segments}, output)
            document = Document(output)

        self.assertEqual(len(document.sections), 2)
        vml_shapes = document.element.body.xpath(
            ".//*[local-name()='shape' and namespace-uri()='urn:schemas-microsoft-com:vml']"
        )
        self.assertEqual(len(document.element.body.xpath(".//wp:anchor")), 0)
        self.assertEqual(len(vml_shapes), 240)
        self.assertEqual(len(document.element.body.xpath(
            ".//*[local-name()='shapetype' and namespace-uri()='urn:schemas-microsoft-com:vml']"
        )), 1)
        self.assertEqual(len({
            shape.get("{urn:schemas-microsoft-com:office:office}spid") for shape in vml_shapes
        }), 240)
        self.assertEqual(len(document.element.body.xpath(".//w:framePr")), 0)
        host_paragraphs = document.element.body.xpath(
            "./w:p[.//*[local-name()='shape' and namespace-uri()='urn:schemas-microsoft-com:vml']]"
        )
        self.assertEqual(len(host_paragraphs), 240)
        self.assertTrue(all(
            len(paragraph.xpath(
                "./w:r/w:pict/*[local-name()='shape' and namespace-uri()='urn:schemas-microsoft-com:vml']"
            )) == 1
            for paragraph in host_paragraphs
        ))


if __name__ == "__main__":
    unittest.main()
