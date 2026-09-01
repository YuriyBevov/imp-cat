#!/usr/bin/env python3
"""Export the browser layout as page-anchored editable Word/VML text boxes."""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Twips
from lxml import etree


PX_TO_TWIPS = 15  # CSS reference pixel at 96 DPI; Word uses 1/1440 inch.
PX_TO_EMU = 9_525  # CSS reference pixel at 96 DPI; DrawingML uses 914400 EMU/inch.
BASE_RELATIVE_HEIGHT = 251_658_240
WPS_GRAPHIC_URI = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
INVALID_XML_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "o": "urn:schemas-microsoft-com:office:office",
    "v": "urn:schemas-microsoft-com:vml",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w10": "urn:schemas-microsoft-com:office:word",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
}

for namespace_prefix, namespace_uri in NS.items():
    etree.register_namespace(namespace_prefix, namespace_uri)


def px_to_twips(value: float) -> int:
    return max(1, round(float(value) * PX_TO_TWIPS))


def px_to_emu(value: float) -> int:
    return max(1, round(float(value) * PX_TO_EMU))


def element(prefix: str, name: str, attrs: dict | None = None) -> etree._Element:
    node = etree.Element(etree.QName(NS[prefix], name))
    for key, value in (attrs or {}).items():
        node.set(key, str(value))
    return node


def word_attr(node: etree._Element, name: str, value: object) -> None:
    node.set(etree.QName(NS["w"], name), str(value))


def namespaced_attr(node: etree._Element, prefix: str, name: str, value: object) -> None:
    node.set(etree.QName(NS[prefix], name), str(value))


def configure_section(section, page: dict) -> None:
    section.page_width = Twips(px_to_twips(page["widthPx"]))
    section.page_height = Twips(px_to_twips(page["heightPx"]))
    section.top_margin = Twips(0)
    section.right_margin = Twips(0)
    section.bottom_margin = Twips(0)
    section.left_margin = Twips(0)
    section.header_distance = Twips(0)
    section.footer_distance = Twips(0)


def create_run_properties(segment: dict) -> etree._Element:
    properties = element("w", "rPr")
    font_name = INVALID_XML_CHARS.sub("", str(segment.get("fontFamily") or "Arial"))[:200]
    fonts = element("w", "rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        word_attr(fonts, attribute, font_name)
    properties.append(fonts)

    if float(segment.get("fontWeight", 400)) >= 600:
        properties.append(element("w", "b"))
        properties.append(element("w", "bCs"))

    half_points = max(2, round(float(segment["fontSizePx"]) * 1.5))
    size = element("w", "sz")
    word_attr(size, "val", half_points)
    complex_size = element("w", "szCs")
    word_attr(complex_size, "val", half_points)
    properties.extend((size, complex_size))

    color = element("w", "color")
    word_attr(color, "val", str(segment.get("color", "#111827")).lstrip("#").upper())
    properties.append(color)
    return properties


def append_text_content(run: etree._Element, text: str) -> None:
    cleaned = INVALID_XML_CHARS.sub("", text)
    for token in re.split(r"(\n|\t)", cleaned):
        if token == "\n":
            run.append(element("w", "br"))
        elif token == "\t":
            run.append(element("w", "tab"))
        elif token:
            text_node = element("w", "t")
            text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            text_node.text = token
            run.append(text_node)
    if not cleaned:
        run.append(element("w", "t"))


def create_text_box_paragraph(segment: dict) -> etree._Element:
    paragraph = element("w", "p")
    paragraph_properties = element("w", "pPr")
    spacing = element("w", "spacing")
    word_attr(spacing, "before", 0)
    word_attr(spacing, "after", 0)
    word_attr(
        spacing,
        "line",
        max(1, round(float(segment["fontSizePx"]) * PX_TO_TWIPS * float(segment.get("lineHeight", 1.2)))),
    )
    word_attr(spacing, "lineRule", "exact")
    paragraph_properties.append(spacing)

    alignment = element("w", "jc")
    word_attr(alignment, "val", segment.get("alignment", "left"))
    paragraph_properties.append(alignment)
    paragraph_properties.append(element("w", "contextualSpacing"))
    paragraph.append(paragraph_properties)

    run = element("w", "r")
    run.append(create_run_properties(segment))
    append_text_content(run, segment["text"])
    paragraph.append(run)
    return paragraph


def create_vml_text_box(
    segment: dict,
    shape_id: int,
    include_shape_type: bool = False,
) -> etree._Element:
    pict = element("w", "pict")
    if include_shape_type:
        shape_type = element(
            "v",
            "shapetype",
            {
                "id": "_x0000_t202",
                "coordsize": "21600,21600",
                "path": "m,l,21600r21600,l21600,xe",
            },
        )
        namespaced_attr(shape_type, "o", "spt", 202)
        shape_type.append(element("v", "stroke", {"joinstyle": "miter"}))
        shape_path = element("v", "path", {"gradientshapeok": "t"})
        namespaced_attr(shape_path, "o", "connecttype", "rect")
        shape_type.append(shape_path)
        pict.append(shape_type)

    points = lambda value: f"{float(value) * 0.75:.3f}".rstrip("0").rstrip(".")
    style = ";".join(
        (
            "position:absolute",
            "left:0",
            "top:0",
            "text-align:left",
            f"margin-left:{points(segment['x'])}pt",
            f"margin-top:{points(segment['y'])}pt",
            f"width:{points(segment['width'])}pt",
            f"height:{points(segment['height'])}pt",
            f"z-index:{BASE_RELATIVE_HEIGHT + int(segment.get('zIndex', shape_id))}",
            "visibility:visible",
            "mso-wrap-style:square",
            "mso-width-percent:0",
            "mso-height-percent:0",
            "mso-wrap-distance-left:0",
            "mso-wrap-distance-top:0",
            "mso-wrap-distance-right:0",
            "mso-wrap-distance-bottom:0",
            "mso-position-horizontal:absolute",
            "mso-position-horizontal-relative:page",
            "mso-position-vertical:absolute",
            "mso-position-vertical-relative:page",
            "mso-width-relative:page",
            "mso-height-relative:page",
            "v-text-anchor:top",
        )
    )
    shape = element(
        "v",
        "shape",
        {
            "id": f"ICAT segment {shape_id}",
            "type": "#_x0000_t202",
            "style": style,
            "filled": "f",
            "stroked": "f",
        },
    )
    namespaced_attr(shape, "o", "spid", f"_x0000_s{1024 + shape_id}")
    namespaced_attr(shape, "o", "allowoverlap", "t")
    namespaced_attr(shape, "o", "allowincell", "f")
    text_box = element("v", "textbox", {"inset": "0,0,0,0"})
    text_box_content = element("w", "txbxContent")
    text_box_content.append(create_text_box_paragraph(segment))
    text_box.append(text_box_content)
    shape.append(text_box)
    pict.append(shape)
    return pict


def create_drawingml_text_box(segment: dict, shape_id: int) -> etree._Element:
    width_emu = px_to_emu(segment["width"])
    height_emu = px_to_emu(segment["height"])
    drawing = element("w", "drawing")
    anchor = element(
        "wp",
        "anchor",
        {
            "distT": "0",
            "distB": "0",
            "distL": "0",
            "distR": "0",
            "simplePos": "0",
            "relativeHeight": str(BASE_RELATIVE_HEIGHT + int(segment.get("zIndex", shape_id))),
            "behindDoc": "0",
            "locked": "0",
            "layoutInCell": "0",
            "allowOverlap": "1",
        },
    )

    anchor.append(element("wp", "simplePos", {"x": "0", "y": "0"}))
    horizontal = element("wp", "positionH", {"relativeFrom": "page"})
    horizontal_offset = element("wp", "posOffset")
    horizontal_offset.text = str(px_to_emu(segment["x"]))
    horizontal.append(horizontal_offset)
    anchor.append(horizontal)

    vertical = element("wp", "positionV", {"relativeFrom": "page"})
    vertical_offset = element("wp", "posOffset")
    vertical_offset.text = str(px_to_emu(segment["y"]))
    vertical.append(vertical_offset)
    anchor.append(vertical)

    anchor.append(element("wp", "extent", {"cx": width_emu, "cy": height_emu}))
    anchor.append(element("wp", "effectExtent", {"l": "0", "t": "0", "r": "0", "b": "0"}))
    anchor.append(element("wp", "wrapNone"))
    anchor.append(
        element(
            "wp",
            "docPr",
            {
                "id": shape_id,
                "name": f"ICAT segment {shape_id}",
                "descr": str(segment.get("id", "segment"))[:200],
            },
        )
    )
    anchor.append(element("wp", "cNvGraphicFramePr"))

    graphic = element("a", "graphic")
    graphic_data = element("a", "graphicData", {"uri": WPS_GRAPHIC_URI})
    shape = element("wps", "wsp")
    shape.append(element("wps", "cNvSpPr", {"txBox": "1"}))

    shape_properties = element("wps", "spPr")
    transform = element("a", "xfrm")
    transform.append(element("a", "off", {"x": "0", "y": "0"}))
    transform.append(element("a", "ext", {"cx": width_emu, "cy": height_emu}))
    shape_properties.append(transform)
    geometry = element("a", "prstGeom", {"prst": "rect"})
    geometry.append(element("a", "avLst"))
    shape_properties.append(geometry)
    shape_properties.append(element("a", "noFill"))
    line = element("a", "ln")
    line.append(element("a", "noFill"))
    shape_properties.append(line)
    shape.append(shape_properties)

    text_box = element("wps", "txbx")
    text_box_content = element("w", "txbxContent")
    text_box_content.append(create_text_box_paragraph(segment))
    text_box.append(text_box_content)
    shape.append(text_box)

    body_properties = element(
        "wps",
        "bodyPr",
        {
            "rot": "0",
            "spcFirstLastPara": "0",
            "vertOverflow": "overflow",
            "horzOverflow": "overflow",
            "vert": "horz",
            "wrap": "square",
            "lIns": "0",
            "tIns": "0",
            "rIns": "0",
            "bIns": "0",
            "numCol": "1",
            "spcCol": "0",
            "rtlCol": "0",
            "fromWordArt": "0",
            "anchor": "t",
            "anchorCtr": "0",
            "forceAA": "0",
            "compatLnSpc": "1",
        },
    )
    text_warp = element("a", "prstTxWarp", {"prst": "textNoShape"})
    text_warp.append(element("a", "avLst"))
    body_properties.append(text_warp)
    body_properties.append(element("a", "noAutofit"))
    shape.append(body_properties)

    graphic_data.append(shape)
    graphic.append(graphic_data)
    anchor.append(graphic)
    drawing.append(anchor)

    alternate_content = element("mc", "AlternateContent")
    choice = element("mc", "Choice", {"Requires": "wps"})
    choice.append(drawing)
    fallback = element("mc", "Fallback")
    fallback.append(create_vml_text_box(segment, shape_id))
    alternate_content.extend((choice, fallback))
    return alternate_content


def create_text_box_drawing(
    segment: dict,
    shape_id: int,
    include_shape_type: bool = False,
) -> etree._Element:
    """Use Transitional VML directly for maximum desktop Word compatibility.

    The DrawingML implementation is retained above for a later opt-in mode, but
    the prototype exporter intentionally avoids Office-version-specific WPS
    markup after Word rejected dense documents containing many WPS shapes.
    """
    return create_vml_text_box(segment, shape_id, include_shape_type)


def add_shape_anchor(
    document: Document,
    segment: dict,
    shape_id: int,
    include_shape_type: bool = False,
) -> None:
    """Put one floating text box in one flow paragraph.

    Desktop Word rejects some dense VML documents when dozens of independent
    ``w:pict`` text boxes are children of the same paragraph, even though the
    ZIP and XML remain valid.  A separate, one-twip host paragraph mirrors the
    structure Word itself writes and isolates every editable shape.
    """
    paragraph = document.add_paragraph()
    paragraph_properties = paragraph._p.get_or_add_pPr()
    spacing = element("w", "spacing")
    word_attr(spacing, "before", 0)
    word_attr(spacing, "after", 0)
    word_attr(spacing, "line", 1)
    word_attr(spacing, "lineRule", "exact")
    paragraph_properties.append(spacing)

    run = element("w", "r")
    run_properties = element("w", "rPr")
    run_properties.append(element("w", "noProof"))
    run_size = element("w", "sz")
    word_attr(run_size, "val", 2)
    run_properties.append(run_size)
    run.append(run_properties)
    run.append(create_text_box_drawing(segment, shape_id, include_shape_type=include_shape_type))
    paragraph._p.append(run)


def add_page_anchors(document: Document, segments: list[dict], first_shape_id: int) -> int:
    shape_id = first_shape_id
    for segment in segments:
        add_shape_anchor(document, segment, shape_id, include_shape_type=shape_id == 1)
        shape_id += 1
    return shape_id


def configure_vml_settings(document: Document, segment_count: int) -> None:
    """Keep Word's declared VML ID range above every generated text box."""
    settings = document.settings.element
    shape_defaults = settings.find(qn("w:shapeDefaults"))
    if shape_defaults is None:
        shape_defaults = element("w", "shapeDefaults")
        settings.append(shape_defaults)

    vml_defaults = shape_defaults.find(f"{{{NS['o']}}}shapedefaults")
    if vml_defaults is None:
        vml_defaults = element("o", "shapedefaults")
        namespaced_attr(vml_defaults, "v", "ext", "edit")
        shape_defaults.insert(0, vml_defaults)
    vml_defaults.set("spidmax", str(max(1027, 1025 + segment_count)))

    for compatibility_setting in settings.xpath("./w:compat/w:compatSetting"):
        if compatibility_setting.get(qn("w:name")) == "compatibilityMode":
            compatibility_setting.set(qn("w:val"), "15")


def validate_export(output_path: Path, expected_pages: int, expected_segments: int) -> None:
    """Fail the HTTP export before download if the package violates Word/VML invariants."""
    try:
        with ZipFile(output_path) as archive:
            if archive.testzip() is not None:
                raise ValueError("DOCX contains a damaged ZIP entry")
            required_parts = {"[Content_Types].xml", "word/document.xml", "word/settings.xml"}
            missing_parts = required_parts.difference(archive.namelist())
            if missing_parts:
                raise ValueError(f"DOCX is missing required parts: {sorted(missing_parts)}")
            for member in archive.namelist():
                if member.endswith((".xml", ".rels")):
                    etree.fromstring(archive.read(member))
    except BadZipFile as error:
        raise ValueError("DOCX is not a valid ZIP package") from error

    reopened = Document(output_path)
    if len(reopened.sections) != expected_pages:
        raise ValueError("DOCX page-section count does not match the layout")

    shapes = reopened.element.body.xpath(
        ".//*[local-name()='shape' and namespace-uri()='urn:schemas-microsoft-com:vml']"
    )
    if len(shapes) != expected_segments:
        raise ValueError("DOCX VML shape count does not match the segment count")
    shape_types = reopened.element.body.xpath(
        ".//*[local-name()='shapetype' and namespace-uri()='urn:schemas-microsoft-com:vml']"
    )
    if expected_segments and len(shape_types) != 1:
        raise ValueError("DOCX must define the VML text-box shape type exactly once")

    host_paragraphs = reopened.element.body.xpath(
        "./w:p[.//*[local-name()='shape' and namespace-uri()='urn:schemas-microsoft-com:vml']]"
    )
    if any(len(paragraph.xpath(
        "./w:r/w:pict/*[local-name()='shape' and namespace-uri()='urn:schemas-microsoft-com:vml']"
    )) != 1 for paragraph in host_paragraphs):
        raise ValueError("Each DOCX VML shape must have an independent host paragraph")
    if len(host_paragraphs) != expected_segments:
        raise ValueError("DOCX VML host-paragraph count does not match the segment count")

    shape_ids = [shape.get(f"{{{NS['o']}}}spid") for shape in shapes]
    if any(not shape_id for shape_id in shape_ids) or len(set(shape_ids)) != len(shape_ids):
        raise ValueError("DOCX contains missing or duplicate VML shape IDs")
    maximum_shape_id = max(
        (int(shape_id.rsplit("s", 1)[-1]) for shape_id in shape_ids),
        default=0,
    )
    shape_defaults = reopened.settings.element.xpath(
        ".//*[local-name()='shapedefaults' and namespace-uri()='urn:schemas-microsoft-com:office:office']"
    )
    declared_maximum = max((int(node.get("spidmax", 0)) for node in shape_defaults), default=0)
    if maximum_shape_id >= declared_maximum:
        raise ValueError("DOCX VML spidmax is lower than an emitted shape ID")


def export_layout(payload: dict, output_path: Path) -> None:
    document = Document()
    document.core_properties.title = payload.get("title", "ICAT Grid export")

    pages = payload["pages"]
    grouped_segments = {
        page["index"]: sorted(
            (segment for segment in payload["segments"] if segment["pageIndex"] == page["index"]),
            key=lambda item: (item.get("zIndex", 0), item["y"], item["x"]),
        )
        for page in pages
    }

    next_shape_id = 1
    for page_index, page in enumerate(pages):
        section = document.sections[0] if page_index == 0 else document.add_section(WD_SECTION.NEW_PAGE)
        configure_section(section, page)
        next_shape_id = add_page_anchors(document, grouped_segments[page["index"]], next_shape_id)

    configure_vml_settings(document, len(payload["segments"]))
    document.save(output_path)
    validate_export(output_path, len(pages), len(payload["segments"]))


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("usage: export_docx.py LAYOUT.json OUTPUT.docx")
    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    payload = json.loads(input_path.read_text(encoding="utf-8"))
    export_layout(payload, output_path)
    print(json.dumps({"pages": len(payload["pages"]), "segments": len(payload["segments"])}))


if __name__ == "__main__":
    main()
