#!/usr/bin/env python3
"""Generate a representative, non-sensitive DOCX for the browser smoke test."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def main() -> None:
    output = Path(__file__).resolve().parents[1] / "fixtures" / "sample.docx"
    output.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    title = document.add_heading("ICAT Grid — тестовый документ", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = document.add_paragraph()
    paragraph.add_run("Этот текст можно ").bold = False
    emphasized = paragraph.add_run("редактировать и перемещать")
    emphasized.bold = True
    paragraph.add_run(" поверх виртуальной сетки.")

    document.add_heading("Гипотеза", level=2)
    document.add_paragraph(
        "Каждый абзац становится сегментом со стабильным ID, координатами страницы, размером и типографикой."
    )
    document.add_paragraph(
        "После изменения положения браузер отправляет координатную модель на сервер, который собирает новый DOCX."
    )

    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Исходный сегмент"
    table.cell(0, 1).text = "Переведённый сегмент"
    table.cell(1, 0).text = "Move me"
    table.cell(1, 1).text = "Перемести меня"

    note = document.add_paragraph("Ограничение прототипа: изображения и сложные Word-объекты пока не экспортируются.")
    note.runs[0].font.size = Pt(9)
    document.save(output)


if __name__ == "__main__":
    main()
